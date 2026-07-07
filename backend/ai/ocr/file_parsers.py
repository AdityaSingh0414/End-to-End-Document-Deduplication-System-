import os
import pypdf
import docx
import logging

logger = logging.getLogger("file_parsers")


def extract_text_from_file(file_path: str) -> str:
    """
    Extract digital text from PDF and DOCX files.
    Returns an empty string if extraction fails or format is unsupported.
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return ""
        
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == ".pdf":
        try:
            text_parts = []
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page_num in range(len(reader.pages)):
                    page_text = reader.pages[page_num].extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return ""
            
    elif ext == ".docx":
        try:
            doc = docx.Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return ""
            
    elif ext in [".txt", ".log", ".json", ".csv"]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""
            
    return ""
